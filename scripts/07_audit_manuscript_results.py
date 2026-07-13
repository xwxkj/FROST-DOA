#!/usr/bin/env python3
"""Audit manuscript tables and names against terminal-generated benchmark CSVs.

The script never edits the manuscript. It produces a machine-readable JSON and
a concise Markdown report. A nonzero exit status means the manuscript still
contains stale metrics or internal development names and must not be submitted.
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

PUBLIC_METHODS = {
    "FROST-DOA",
    "MUSIC + MDL",
    "Root-MUSIC + MDL",
    "ESPRIT + MDL",
    "FBSS-MUSIC + MDL",
    "Toeplitz-MUSIC + MDL",
    "SOMP + MDL",
    "SBL + MDL",
    "DA-MUSIC + MDL",
    "SubspaceNet + residual selector",
}
FORBIDDEN_PATTERNS = [
    r"FROST[-–— ]?DOA\s+v\d+",
    r"FROST[-–— ]?V\d+",
    r"Oracle[- ]K",
    r"debug run",
]
ALIASES = {
    "FROST-DOA v16": "FROST-DOA",
    "DA-MUSIC+MDL": "DA-MUSIC + MDL",
    "SubspaceNet+Residual-K": "SubspaceNet + residual selector",
    "SubspaceNet + Residual-K": "SubspaceNet + residual selector",
}


def normalize_method(text: str) -> str:
    value = " ".join(text.replace("\u2013", "-").replace("\u2014", "-").split())
    return ALIASES.get(value, value)


def parse_number(text: str) -> float | None:
    cleaned = text.replace("%", "").replace("°", "").replace(",", "").strip()
    match = re.search(r"[-+]?\d+(?:\.\d+)?", cleaned)
    return float(match.group()) if match else None


def document_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def find_method_rows(document: Document) -> dict[str, list[str]]:
    found: dict[str, list[str]] = {}
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if not cells:
                continue
            method = normalize_method(cells[0])
            if method in PUBLIC_METHODS:
                found[method] = cells
    return found


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manuscript", type=Path, required=True)
    parser.add_argument("--overall-summary", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, default=Path("audit/manuscript"))
    parser.add_argument("--tolerance", type=float, default=0.005)
    args = parser.parse_args()

    document = Document(args.manuscript)
    overall = pd.read_csv(args.overall_summary).set_index("method")
    rows = find_method_rows(document)
    full_text = document_text(document)
    issues: list[dict[str, Any]] = []

    for pattern in FORBIDDEN_PATTERNS:
        matches = sorted(set(re.findall(pattern, full_text, flags=re.IGNORECASE)))
        if matches:
            issues.append(
                {"type": "forbidden_internal_name", "pattern": pattern, "matches": matches}
            )

    comparisons = []
    for method, expected in overall.iterrows():
        if method not in rows:
            issues.append({"type": "method_missing_from_manuscript_table", "method": method})
            continue
        cells = rows[method]
        numbers = [
            value for value in (parse_number(cell) for cell in cells[1:]) if value is not None
        ]
        expected_values = [
            float(expected["mean_rmse_deg"]),
            float(expected["median_rmse_deg"]),
            float(expected["source_number_accuracy_pct"]),
        ]
        labels = ["mean_rmse_deg", "median_rmse_deg", "source_number_accuracy_pct"]
        if len(numbers) < 3:
            issues.append({"type": "insufficient_numeric_cells", "method": method, "cells": cells})
            continue
        for label, manuscript_value, terminal_value in zip(labels, numbers[:3], expected_values):
            difference = manuscript_value - terminal_value
            comparisons.append(
                {
                    "method": method,
                    "metric": label,
                    "manuscript_value": manuscript_value,
                    "terminal_value": terminal_value,
                    "difference": difference,
                }
            )
            if abs(difference) > args.tolerance:
                issues.append(
                    {
                        "type": "stale_metric",
                        "method": method,
                        "metric": label,
                        "manuscript_value": manuscript_value,
                        "terminal_value": terminal_value,
                        "difference": difference,
                    }
                )

    result = {
        "status": "pass" if not issues else "fail",
        "manuscript": str(args.manuscript),
        "overall_summary": str(args.overall_summary),
        "tolerance": args.tolerance,
        "issues": issues,
        "comparisons": comparisons,
    }
    args.output_dir.mkdir(parents=True, exist_ok=True)
    (args.output_dir / "manuscript_result_audit.json").write_text(
        json.dumps(result, indent=2), encoding="utf-8"
    )
    report = [
        "# Manuscript/result consistency audit",
        "",
        f"- Status: **{result['status'].upper()}**",
        f"- Manuscript: `{args.manuscript}`",
        f"- Terminal summary: `{args.overall_summary}`",
        f"- Numeric tolerance: `{args.tolerance}`",
        "",
    ]
    if issues:
        report += ["## Required corrections", ""]
        for issue in issues:
            report.append(f"- `{issue['type']}`: {json.dumps(issue, ensure_ascii=False)}")
    else:
        report.append("No stale metrics or internal development names were detected.")
    (args.output_dir / "manuscript_result_audit.md").write_text("\n".join(report), encoding="utf-8")
    print("\n".join(report))
    raise SystemExit(0 if not issues else 2)


if __name__ == "__main__":
    main()
